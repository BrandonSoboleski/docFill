#//////////////////////////////////////////////////////
#Imports                                              /
#//////////////////////////////////////////////////////
from docx import Document
import os
#from openpyxl import load_workbook
#//////////////////////////////////////////////////////
#Docx Files                                           /
#//////////////////////////////////////////////////////

def editmtf(patientName, adDate, Diagnosis):
    print("EDITING MTF FORM")
    maind = os.getcwd()
    os.chdir('copiedFrom')
    mtf = Document("medicare tracking form - Copy.docx")
    for paragraph in mtf.paragraphs:
        if '______________________________________________' in paragraph.text:
            paragraph.text = paragraph.text.split("______________________________________________")[0], patientName, paragraph.text.split("______________________________________________")[1].split("_____/_____/______")[0], " ", adDate
            print(paragraph.text)
        elif '______________________________________' in paragraph.text:
            paragraph.text = paragraph.text.split("______________________________________")[0], Diagnosis
            print(paragraph.text)
    os.chdir(maind)
    os.chdir('output')
    mtf.save("medicare tracking form - Copy.docx")
    os.chdir(maind)
    
def editrf3(residentName):
    print("EDITING RF3 FORM")
    maind = os.getcwd()
    os.chdir('copiedFrom')
    rf3 = Document("REVIEW FORM 3.0 computerized - Copy.docx")
    for paragraph in rf3.paragraphs:
        if '                                           ' in paragraph.text:
            paragraph.text = paragraph.text.split("                                         ")[0], residentName, " ", paragraph.text.split("                                         ")[1]
            print(paragraph.text)
    os.chdir(maind)
    os.chdir('output')
    rf3.save("REVIEW FORM 3.0 computerized - Copy.docx")
    os.chdir(maind)
    
def editmmd(residentName, adDate, Diagnosis):
    print("EDITING MMD FORM")
    maind = os.getcwd()
    os.chdir('copiedFrom')
    mmd = Document("Medicare Meeting Documentation - Copy.docx")

    for table in mmd.tables:
        for row in table.rows:
            for cell in row.cells:
                if 'Resident Name:' in cell.text:
                    cell.text = cell.text.split("Resident Name:")[0], "Resident Name: ", residentName, " ", cell.text.split("Resident Name:")[1]
                    print(cell.text)
                elif 'Admission Date:' in cell.text:
                    cell.text = cell.text.split("Admission Date:")[0], "Admission Date: ", adDate, " ", cell.text.split("Admission Date:")[1]
                    print(cell.text)
                elif 'Diagnosis:' in cell.text:
                    cell.text = cell.text.split("Diagnosis:")[0], "Diagnosis: ", Diagnosis, " ", cell.text.split("Diagnosis:")[1]
                    print(cell.text)

    os.chdir(maind)
    os.chdir('output')
    mmd.save("Medicare Meeting Documentation - Copy.docx")
    os.chdir(maind)
    
#//////////////////////////////////////////////////////
#Doc File       Unused                                /
#//////////////////////////////////////////////////////

#def edit():
#    ole = olefile.OleFileIO('Medicare Meeting Documentation - Copy.doc', write_mode=True)
#    data = ole.openstream('Medicare Meeting Documentation - Copy.doc').read()
#    #data = data.replace(b'Resident Name:', b'Resident Name: test')
#    ole.write_stream('Medicare Meeting Documentation - Copy.doc', data)
#    ole.close()

#//////////////////////////////////////////////////////
#Excel Files                                          /
#//////////////////////////////////////////////////////
def excelEdit():
    wb = load_workbook('Med A tracking form - Copy.xlsx')
    ws = wb['Sheet1']
    ws['A1'] = 'A1'
    wb.save('Med A tracking form - Copy.xlsx')    
#//////////////////////////////////////////////////////
#Function Calls                                       /
#//////////////////////////////////////////////////////

def editall(args):
    editmtf(args[0], args[1], args[2])
    editmmd(args[0], args[1], args[2])
    editrf3(args[0])


#Testing   
#editmtf("Gloria", "2/2/2022", "Has the Flu")
#editrf3("Gloria")
#editmmd("Gloria", "2/2/2022", "Has the Flu")
#excelEdit()
#def test():
#    print("success")
